import os
import logging
from jinja2 import Environment, FileSystemLoader
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except Exception as e:
    logger = logging.getLogger(__name__)
    logger.warning(f"WeasyPrint could not be loaded ({e}). PDF rendering will fall back to a minimal PDF generator.")
    WEASYPRINT_AVAILABLE = False

from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)

# Map theme names to hex colors
THEME_COLORS = {
    "default": "#111827",  # slate/gray
    "navy": "#1A365D",     # dark blue
    "emerald": "#047857",  # green
    "indigo": "#4338CA",   # indigo
    "violet": "#6D28D9",   # purple
    "rose": "#BE123C"      # rose red
}

class ResumeRenderer:
    def __init__(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        templates_path = os.path.join(current_dir, "../templates/resume")
        self.env = Environment(loader=FileSystemLoader(templates_path))

    async def render_pdf(self, resume_data: dict, template_name: str, color_theme: str = "default") -> bytes:
        """Renders the resume data to PDF using WeasyPrint (or fallback if unavailable)."""
        logger.info(f"Rendering PDF resume with template: {template_name}, theme: {color_theme}...")
        
        if not WEASYPRINT_AVAILABLE:
            logger.warning("WeasyPrint is not available. Generating a minimal fallback PDF.")
            name = resume_data.get("name", "Candidate")
            email = resume_data.get("email", "")
            summary = resume_data.get("summary", "")
            
            stream_content = f"BT\n/F1 18 Tf\n50 700 Td\n({name} - Resume) Tj\n0 -30 Td\n/F1 12 Tf\n({email}) Tj\n0 -40 Td\n(Summary:) Tj\n0 -20 Td\n({summary[:60]}...) Tj\n0 -40 Td\n(Note: Install Pango/Glib on your host for high quality WeasyPrint PDF layout.) Tj\nET"
            stream_bytes = stream_content.encode("utf-8")
            stream_len = len(stream_bytes)
            
            fallback_pdf_parts = [
                b"%PDF-1.4",
                b"1 0 obj",
                b"<< /Type /Catalog /Pages 2 0 R >>",
                b"endobj",
                b"2 0 obj",
                b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
                b"endobj",
                b"3 0 obj",
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
                b"endobj",
                b"4 0 obj",
                b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
                b"endobj",
                b"5 0 obj",
                f"<< /Length {stream_len} >>".encode("utf-8"),
                b"stream",
                stream_bytes,
                b"endstream",
                b"endobj",
                b"xref",
                b"0 6",
                b"0000000000 65535 f ",
                b"0000000009 00000 n ",
                b"0000000058 00000 n ",
                b"0000000115 00000 n ",
                b"0000000223 00000 n ",
                b"0000000290 00000 n ",
                b"trailer",
                b"<< /Size 6 /Root 1 0 R >>",
                b"startxref",
                b"385",
                b"%%EOF"
            ]
            return b"\n".join(fallback_pdf_parts)

        try:
            # 1. Resolve templates and variables
            template_file = f"{template_name}.html"
            template = self.env.get_template(template_file)
            
            theme_color = THEME_COLORS.get(color_theme, THEME_COLORS["default"])
            
            # 2. Compile HTML with Jinja2 context
            html_content = template.render(
                resume=resume_data,
                theme_color=theme_color
            )
            
            # 3. Compile to PDF bytes in-memory
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except Exception as e:
            logger.error(f"Failed to render PDF: {e}")
            raise

    async def render_docx(self, resume_data: dict) -> bytes:
        """Renders a simple ATS-safe single column DOCX file using python-docx."""
        logger.info("Rendering DOCX resume...")
        try:
            doc = Document()
            
            # Set default styles
            doc.styles['Normal'].font.name = 'Arial'
            
            # Header
            name = resume_data.get("name", "Candidate")
            doc.add_heading(name, level=0)
            
            contact_parts = []
            if resume_data.get("email"): contact_parts.append(resume_data["email"])
            if resume_data.get("phone"): contact_parts.append(resume_data["phone"])
            if resume_data.get("location"): contact_parts.append(resume_data["location"])
            
            contact_str = " | ".join(contact_parts)
            doc.add_paragraph(contact_str)
            
            links = []
            if resume_data.get("linkedin"): links.append(f"LinkedIn: {resume_data['linkedin']}")
            if resume_data.get("github"): links.append(f"GitHub: {resume_data['github']}")
            if links:
                doc.add_paragraph(" | ".join(links))
                
            # Summary
            summary = resume_data.get("summary")
            if summary:
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(summary)
                
            # Experience
            experience = resume_data.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=1)
                for job in experience:
                    p = doc.add_paragraph()
                    company = job.get("company", "Company")
                    title = job.get("title", "Role")
                    duration = job.get("duration", "")
                    p.add_run(f"{company} - {title}").bold = True
                    p.add_run(f" ({duration})").italic = True
                    
                    for bullet in job.get("bullets", []):
                        doc.add_paragraph(bullet, style='List Bullet')
                        
            # Projects
            projects = resume_data.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=1)
                for proj in projects:
                    p = doc.add_paragraph()
                    proj_name = proj.get("name", "Project")
                    tech = ", ".join(proj.get("tech_stack", []))
                    p.add_run(f"{proj_name} ({tech})").bold = True
                    doc.add_paragraph(proj.get("description", ""))
                    
            # Skills
            skills = resume_data.get("skills", {})
            if skills:
                doc.add_heading("Skills", level=1)
                if skills.get("technical"):
                    doc.add_paragraph(f"Technical Skills: {', '.join(skills['technical'])}")
                if skills.get("soft"):
                    doc.add_paragraph(f"Soft Skills: {', '.join(skills['soft'])}")
                    
            # Education
            education = resume_data.get("education", [])
            if education:
                doc.add_heading("Education", level=1)
                for edu in education:
                    p = doc.add_paragraph()
                    inst = edu.get("institution", "")
                    deg = edu.get("degree", "")
                    yr = edu.get("year", "")
                    gpa = f" (GPA: {edu['gpa']})" if edu.get("gpa") else ""
                    p.add_run(f"{inst} - {deg}{gpa}").bold = True
                    p.add_run(f" ({yr})").italic = True
                    
            # Save docx to bytes buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Failed to render DOCX: {e}")
            raise

resume_renderer = ResumeRenderer()
